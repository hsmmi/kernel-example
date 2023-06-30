from docx import Document

# from docx.shared import Inches

document = Document()

document.add_heading("Document Title", 0)

p = document.add_paragraph("A plain paragraph having some ")
p.add_run("bold").bold = True
p.add_run(" and some ")
p.add_run("italic.").italic = True

document.add_heading("Heading, level 1", level=1)
document.add_paragraph("Intense quote", style="Intense Quote")

document.add_paragraph("first item in unordered list", style="List Bullet")
document.add_paragraph("first item in ordered list", style="List Number")
document.add_paragraph("first item in ordered list", style="List Number")

# document.add_picture("monty-truth.png", width=Inches(1.25))

records = (
    (3, "101", "Spam"),
    (7, "422", "Eggs"),
    (4, "631", "Spam, spam, eggs, and spam"),
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Qty"
hdr_cells[1].text = "Id"
hdr_cells[2].text = "Desc"
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc


class word:
    def __init__(self) -> None:
        self.document = Document()

    def add_heading(self, text: str, level: int = 1):
        self.document.add_heading(text, level)

    def add_paragraph(self, text: str, style: str = None):
        self.document.add_paragraph(text, style)

    def add_picture(self, path: str, width: int = None, height: int = None):
        self.document.add_picture(path, width, height)

    def save(self, path: str):
        self.document.save(path)
